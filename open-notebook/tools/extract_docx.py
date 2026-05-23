import os
import sys

from docx import Document


def extract_docx_to_markdown(docx_path):
    if not os.path.exists(docx_path):
        return

    try:
        doc = Document(docx_path)
    except Exception as e:
        return

    markdown_content = []

    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Simple style mapping (can be expanded)
        style_name = para.style.name.lower()

        prefix = ""
        if "heading 1" in style_name:
            prefix = "# "
        elif "heading 2" in style_name:
            prefix = "## "
        elif "heading 3" in style_name:
            prefix = "### "
        elif "heading 4" in style_name:
            prefix = "#### "
        elif "list" in style_name:
            prefix = "- "

        # Handle bold/italic runs (basic implementation)
        formatted_text = ""
        for run in para.runs:
            run_text = run.text
            if not run_text:
                continue

            if run.bold:
                run_text = f"**{run_text}**"
            if run.italic:
                run_text = f"_{run_text}_"
            formatted_text += run_text

        markdown_content.append(f"{prefix}{formatted_text}")

    # Process tables (append to end or interleave if possible - python-docx separates them)
    # Note: python-docx 'paragraphs' doesn't strictly interleave with 'tables' in the document.paragraphs list logic
    # To strictly preserve order, we need to iterate over document.element.body

    # Re-doing extraction to preserve order of paragraphs and tables?
    # For now, let's stick to paragraphs then tables, OR use iter_block_items if available/custom.
    # Actually, let's try a simpler approach first: just paragraphs.
    # If the user needs tables, we can add them.
    # Wait, the user manual usually has tables. checking if tables exist.

    if doc.tables:
        markdown_content.append("\n\n---\n### Extracted Tables\n")
        for table in doc.tables:
            # simple key-value extraction or markdown table creation
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(f"| {' | '.join(cells)} |")

            if rows:
                # Add header separator
                header_sep = f"| {' | '.join(['---'] * len(table.rows[0].cells))} |"
                rows.insert(1, header_sep)
                markdown_content.extend(rows)
                markdown_content.append("")  # newline

    return "\n\n".join(markdown_content)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit(1)

    docx_path = sys.argv[1]
    content = extract_docx_to_markdown(docx_path)
    if content:
        pass
